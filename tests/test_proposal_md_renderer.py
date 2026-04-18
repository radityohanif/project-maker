from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

from proposal_maker.core.parser import parse_file
from proposal_maker.core.renderer import render

EXAMPLE_MD = Path(__file__).resolve().parents[1] / "examples" / "proposal.md"
REFERENCE_MD = Path(__file__).resolve().parents[1] / "reference" / "current-proposal.md"


def _embedded_images(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as zf:
        return [n for n in zf.namelist() if n.startswith("word/media/")]


def test_example_md_renders_to_openable_docx(tmp_path: Path) -> None:
    spec = parse_file(EXAMPLE_MD, image_cache_dir=tmp_path / "imgs")
    out = tmp_path / "example.docx"
    render(spec, out)
    assert out.exists()
    doc = Document(str(out))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Project X — Proposal" in all_text
    assert "Executive Summary" in all_text
    assert "Technology Stack" in all_text


def test_example_md_renders_table_with_expected_dimensions(tmp_path: Path) -> None:
    spec = parse_file(EXAMPLE_MD, image_cache_dir=tmp_path / "imgs")
    out = tmp_path / "example.docx"
    render(spec, out)
    doc = Document(str(out))
    assert doc.tables
    tech_table = next(
        (t for t in doc.tables if any("Backend" in c.text for c in t.rows[1].cells)),
        None,
    )
    assert tech_table is not None
    assert len(tech_table.rows) == 5
    assert len(tech_table.columns) == 3


@pytest.mark.skipif(not REFERENCE_MD.exists(), reason="reference proposal not present")
def test_reference_md_embeds_all_28_images(tmp_path: Path) -> None:
    spec = parse_file(REFERENCE_MD, image_cache_dir=tmp_path / "imgs")
    out = tmp_path / "ref.docx"
    render(spec, out)
    assert len(_embedded_images(out)) == 28


def test_heading_numbering_is_applied(tmp_path: Path) -> None:
    spec = parse_file(EXAMPLE_MD, image_cache_dir=tmp_path / "imgs")
    out = tmp_path / "example.docx"
    render(spec, out)
    doc = Document(str(out))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    prefixed = [h for h in headings if h and h[0].isdigit()]
    assert prefixed, f"Expected numbered headings, got {headings!r}"


def test_toc_field_inserted_when_enabled(tmp_path: Path) -> None:
    spec = parse_file(EXAMPLE_MD, image_cache_dir=tmp_path / "imgs")
    out = tmp_path / "example.docx"
    render(spec, out)
    with zipfile.ZipFile(out) as zf:
        body = zf.read("word/document.xml").decode("utf-8")
    assert "TOC " in body, "Expected a TOC Word field in the body XML."
