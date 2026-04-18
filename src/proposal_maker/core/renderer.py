from __future__ import annotations

import hashlib
import tempfile
import urllib.request
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree

from proposal_maker.core.docx_ext import add_field, add_hyperlink_run
from proposal_maker.core.mermaid import (
    MermaidUnavailableError,
    render_mermaid,
)
from proposal_maker.core.mermaid import (
    is_available as mermaid_available,
)
from proposal_maker.core.models import (
    CodeBlock,
    ImageBlock,
    InlineRun,
    ListBlock,
    MermaidBlock,
    PageBreakBlock,
    ParagraphBlock,
    ProposalSpec,
    QuoteBlock,
    Section,
    TableBlock,
)
from proposal_maker.core.theme import Theme, apply_theme
from shared.schemas.common import References
from shared.utils.files import ensure_parent


class RenderWarnings:
    """Collects non-fatal warnings surfaced during rendering."""

    def __init__(self) -> None:
        self.messages: list[str] = []

    def add(self, msg: str) -> None:
        self.messages.append(msg)


class _HeadingCounter:
    """Track per-level heading counters to produce ``1.1.1`` numbering."""

    def __init__(self, max_level: int = 4) -> None:
        self.max_level = max_level
        self.counts: list[int] = [0] * max_level

    def next_label(self, level: int) -> str:
        if level < 1 or level > self.max_level:
            return ""
        self.counts[level - 1] += 1
        for i in range(level, self.max_level):
            self.counts[i] = 0
        parts = [str(c) for c in self.counts[:level] if c > 0]
        return ".".join(parts)


def render(
    spec: ProposalSpec,
    path: Path,
    references: References | None = None,
    *,
    template_override: Path | None = None,
    theme_override: Path | None = None,
    allow_network: bool = False,
) -> Path:
    """Render ``spec`` into a ``.docx`` file at ``path``.

    ``references`` optionally provides substitution values exposed to paragraph
    text as ``{{ timeline_xlsx }}`` / ``{{ quote_xlsx }}`` placeholders.
    ``template_override``/``theme_override`` take precedence over ``spec.template``.
    """
    template_path = template_override or spec.template.docx_template
    theme_path = theme_override or spec.template.theme

    doc = Document(str(template_path)) if template_path else Document()
    theme = Theme.load(theme_path) if theme_path else Theme()
    apply_theme(doc, theme)

    warnings = RenderWarnings()
    substitutions = (references or References()).as_substitutions()
    tmp_dir = Path(tempfile.mkdtemp(prefix="proposal-maker-"))

    _write_header_logos(doc, spec)
    _write_cover(doc, spec)

    if spec.toc.enabled:
        _write_toc(doc, spec)

    counter = _HeadingCounter(max_level=spec.numbering.max_level)

    for section in spec.sections:
        _write_section(
            doc,
            section,
            substitutions=substitutions,
            warnings=warnings,
            tmp_dir=tmp_dir,
            allow_network=allow_network,
            counter=counter,
            number_headings=spec.numbering.enabled,
            numbering_max_level=spec.numbering.max_level,
        )

    if spec.footer.enabled:
        _write_footer(doc, spec)

    if warnings.messages:
        doc.add_paragraph()
        note = doc.add_paragraph()
        run = note.add_run("Warnings during generation:")
        run.bold = True
        for msg in warnings.messages:
            doc.add_paragraph(msg, style="List Bullet")

    ensure_parent(path)
    doc.save(path)
    return path


def _write_header_logos(doc: DocxDocument, spec: ProposalSpec) -> None:
    if not spec.logos:
        return
    header = doc.sections[0].header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    for logo in spec.logos:
        if not Path(logo.path).exists():
            continue
        run = para.add_run()
        run.add_picture(str(logo.path), width=Cm(logo.width_cm))
        para.add_run("  ")


def _write_cover(doc: DocxDocument, spec: ProposalSpec) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(spec.meta.name)
    title_run.bold = True
    title_run.font.size = Pt(26)

    if spec.meta.subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(spec.meta.subtitle)
        run.italic = True
        run.font.size = Pt(14)

    meta_lines: list[str] = []
    if spec.meta.client:
        meta_lines.append(f"Client: {spec.meta.client}")
    if spec.meta.author:
        meta_lines.append(f"Prepared by: {spec.meta.author}")
    if spec.meta.date:
        meta_lines.append(f"Date: {spec.meta.date}")
    if spec.meta.version:
        meta_lines.append(f"Version: {spec.meta.version}")
    if spec.meta.doc_id:
        meta_lines.append(f"Doc ID: {spec.meta.doc_id}")
    if spec.meta.confidential:
        meta_lines.append("CONFIDENTIAL")
    for line in meta_lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        run.font.size = Pt(11)


def _write_toc(doc: DocxDocument, spec: ProposalSpec) -> None:
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(spec.toc.title)
    title_run.bold = True
    title_run.font.size = Pt(16)

    toc_para = doc.add_paragraph()
    add_field(toc_para, f'TOC \\o "1-{spec.toc.depth}" \\h \\z \\u')

    page_para = doc.add_paragraph()
    page_para.add_run().add_break(WD_BREAK.PAGE)


def _write_section(
    doc: DocxDocument,
    section: Section,
    *,
    substitutions: dict[str, str],
    warnings: RenderWarnings,
    tmp_dir: Path,
    allow_network: bool,
    counter: _HeadingCounter,
    number_headings: bool,
    numbering_max_level: int,
) -> None:
    heading_text = section.heading
    if number_headings and section.level <= numbering_max_level:
        label = counter.next_label(section.level)
        if label and not _looks_prenumbered(section.heading):
            heading_text = f"{label}. {section.heading}"
    _add_heading(doc, heading_text, section.level)

    for block in section.blocks:
        _write_block(
            doc,
            block,
            substitutions=substitutions,
            warnings=warnings,
            tmp_dir=tmp_dir,
            allow_network=allow_network,
        )
    for child in section.sections:
        _write_section(
            doc,
            child,
            substitutions=substitutions,
            warnings=warnings,
            tmp_dir=tmp_dir,
            allow_network=allow_network,
            counter=counter,
            number_headings=number_headings,
            numbering_max_level=numbering_max_level,
        )


def _add_heading(doc: DocxDocument, text: str, level: int) -> None:
    if 1 <= level <= 4:
        doc.add_heading(text, level=level)
    elif level >= 5:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
    else:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True


def _write_block(
    doc: DocxDocument,
    block: ParagraphBlock
    | ListBlock
    | MermaidBlock
    | ImageBlock
    | PageBreakBlock
    | TableBlock
    | QuoteBlock
    | CodeBlock,
    *,
    substitutions: dict[str, str],
    warnings: RenderWarnings,
    tmp_dir: Path,
    allow_network: bool,
) -> None:
    if isinstance(block, ParagraphBlock):
        para = doc.add_paragraph()
        _write_runs(para, _paragraph_runs(block), substitutions)
        return
    if isinstance(block, ListBlock):
        style = "List Number" if block.ordered else "List Bullet"
        for runs in block.iter_item_runs():
            para = doc.add_paragraph(style=style)
            _write_runs(para, runs, substitutions)
        return
    if isinstance(block, TableBlock):
        _write_table(doc, block, substitutions)
        return
    if isinstance(block, QuoteBlock):
        para = doc.add_paragraph(style="Intense Quote") if _has_style(doc, "Intense Quote") else (
            doc.add_paragraph(style="Quote") if _has_style(doc, "Quote") else doc.add_paragraph()
        )
        runs = block.runs or [InlineRun(text=block.text)]
        _write_runs(para, runs, substitutions)
        return
    if isinstance(block, CodeBlock):
        _write_code(doc, block)
        return
    if isinstance(block, MermaidBlock):
        _write_mermaid(doc, block, warnings, tmp_dir)
        return
    if isinstance(block, ImageBlock):
        _write_image(doc, block, warnings, tmp_dir, allow_network)
        return
    if isinstance(block, PageBreakBlock):
        para = doc.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)
        return


def _paragraph_runs(block: ParagraphBlock) -> list[InlineRun]:
    if block.runs:
        return list(block.runs)
    return [InlineRun(text=block.text)]


def _write_runs(
    para,
    runs: list[InlineRun],
    substitutions: dict[str, str],
) -> None:
    for run in runs:
        text = _substitute(run.text, substitutions)
        if run.link_url:
            add_hyperlink_run(para, text, run.link_url)
            continue
        r = para.add_run(text)
        if run.bold:
            r.bold = True
        if run.italic:
            r.italic = True
        if run.underline:
            r.underline = True
        if run.strike:
            r.font.strike = True
        if run.code:
            r.font.name = "Consolas"


def _write_table(doc: DocxDocument, block: TableBlock, substitutions: dict[str, str]) -> None:
    widths = [len(block.header)] + [len(row) for row in block.rows]
    cols = max(widths) if widths else 1
    if cols <= 0:
        return
    rows_total = (1 if block.header else 0) + len(block.rows)
    if rows_total == 0:
        return
    table = doc.add_table(rows=rows_total, cols=cols)
    try:
        table.style = "Light List Accent 1"
    except KeyError:
        pass
    row_idx = 0
    if block.header:
        for col_idx in range(cols):
            cell = table.rows[row_idx].cells[col_idx]
            cell.paragraphs[0].text = ""
            cell_runs = block.header[col_idx] if col_idx < len(block.header) else []
            para = cell.paragraphs[0]
            _write_runs(para, cell_runs, substitutions)
            for run in para.runs:
                run.bold = True
        row_idx += 1
    for row in block.rows:
        for col_idx in range(cols):
            cell = table.rows[row_idx].cells[col_idx]
            cell.paragraphs[0].text = ""
            cell_runs = row[col_idx] if col_idx < len(row) else []
            _write_runs(cell.paragraphs[0], cell_runs, substitutions)
        row_idx += 1
    if block.caption:
        _add_caption(doc, block.caption)


def _write_code(doc: DocxDocument, block: CodeBlock) -> None:
    para = doc.add_paragraph()
    run = para.add_run(block.source)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    if block.caption:
        _add_caption(doc, block.caption)


def _has_style(doc: DocxDocument, name: str) -> bool:
    try:
        _ = doc.styles[name]
        return True
    except KeyError:
        return False


def _write_mermaid(
    doc: DocxDocument,
    block: MermaidBlock,
    warnings: RenderWarnings,
    tmp_dir: Path,
) -> None:
    if not mermaid_available():
        warnings.add(
            "mermaid-cli (mmdc) not found on PATH; mermaid diagrams rendered as text. "
            "Install with `npm install -g @mermaid-js/mermaid-cli`."
        )
        _write_mermaid_fallback(doc, block)
        return
    out_png = tmp_dir / f"mermaid-{abs(hash(block.source))}.png"
    try:
        render_mermaid(block.source, out_png)
    except (MermaidUnavailableError, RuntimeError) as exc:
        warnings.add(f"Mermaid rendering failed: {exc}")
        _write_mermaid_fallback(doc, block)
        return
    doc.add_picture(str(out_png), width=Cm(14))
    if block.caption:
        _add_caption(doc, block.caption)


def _write_mermaid_fallback(doc: DocxDocument, block: MermaidBlock) -> None:
    fence = doc.add_paragraph("```mermaid")
    fence.runs[0].font.name = "Courier New"
    for line in block.source.splitlines() or [block.source]:
        para = doc.add_paragraph(line)
        if para.runs:
            para.runs[0].font.name = "Courier New"
    close = doc.add_paragraph("```")
    close.runs[0].font.name = "Courier New"
    if block.caption:
        _add_caption(doc, block.caption)


def _write_image(
    doc: DocxDocument,
    block: ImageBlock,
    warnings: RenderWarnings,
    tmp_dir: Path,
    allow_network: bool,
) -> None:
    img_path = _resolve_image_source(block, warnings, tmp_dir, allow_network)
    if img_path is None:
        return
    try:
        picture_para = doc.add_paragraph()
        if block.align == "center":
            picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block.align == "right":
            picture_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = picture_para.add_run()
        run.add_picture(str(img_path), width=Cm(block.width_cm))
    except Exception as exc:
        warnings.add(f"Failed to embed image {img_path}: {exc}")
        return
    if block.caption:
        _add_caption(doc, block.caption)


def _resolve_image_source(
    block: ImageBlock,
    warnings: RenderWarnings,
    tmp_dir: Path,
    allow_network: bool,
) -> Path | None:
    if block.path is not None:
        path = Path(block.path)
        if not path.exists():
            warnings.add(f"Image not found, skipped: {path}")
            return None
        return path
    if block.data_uri is not None:
        from proposal_maker.core.md_images import _decode_data_uri

        decoded = _decode_data_uri(block.data_uri, tmp_dir)
        if decoded is None:
            warnings.add("Unparseable data URI for image, skipped.")
        return decoded
    if block.url is not None:
        if not allow_network:
            warnings.add(
                f"Remote image skipped (pass --allow-network to fetch): {block.url}"
            )
            return None
        try:
            req = urllib.request.Request(
                block.url, headers={"User-Agent": "proposal-maker/0.1"}
            )
            with urllib.request.urlopen(req, timeout=15) as resp:  # noqa: S310 - opt-in flag
                data = resp.read()
            suffix = Path(block.url).suffix or ".bin"
            digest = hashlib.sha1(data, usedforsecurity=False).hexdigest()[:16]
            out = tmp_dir / f"remote-{digest}{suffix}"
            out.write_bytes(data)
            return out
        except Exception as exc:
            warnings.add(f"Failed to download {block.url}: {exc}")
            return None
    return None


def _add_caption(doc: DocxDocument, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def _substitute(text: str, substitutions: dict[str, str]) -> str:
    out = text
    for key, value in substitutions.items():
        out = out.replace(f"{{{{ {key} }}}}", value).replace(f"{{{{{key}}}}}", value)
    return out


def _looks_prenumbered(heading: str) -> bool:
    """Return True if the heading already starts with ``N.`` or ``N.M``."""
    import re

    return bool(re.match(r"^\s*\d+(\.\d+)*[.)]?\s+\S", heading))


def _write_footer(doc: DocxDocument, spec: ProposalSpec) -> None:
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for child in list(para._p):
            para._p.remove(child)
        etree.SubElement(para._p, qn("w:pPr"))

        left_bits: list[str] = []
        if spec.footer.text:
            left_bits.append(spec.footer.text)
        if spec.meta.confidential:
            left_bits.append("CONFIDENTIAL")
        prefix = " | ".join(left_bits)
        if prefix:
            run = para.add_run(prefix + "  |  ")
            run.font.size = Pt(9)

        if spec.footer.page_numbers:
            run = para.add_run("Page ")
            run.font.size = Pt(9)
            add_field(para, "PAGE")
            run = para.add_run(" of ")
            run.font.size = Pt(9)
            add_field(para, "NUMPAGES")
